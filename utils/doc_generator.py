from typing import Dict, Any, List, Optional, Tuple

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from utils.kyc_rules import REQUIRED_HEADINGS


def set_paragraph_style(p, size: int = 11):
    for run in p.runs:
        run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15


def validate_aml_narrative(aml_narrative: str) -> None:
    raw_lines = aml_narrative.split("\n")
    lines = [ln.rstrip() for ln in raw_lines if ln.strip()]
    if not lines:
        raise ValueError("Narrative is empty.")

    first = lines[0].strip()
    if first != REQUIRED_HEADINGS[0]:
        raise ValueError(
            f"Narrative must start with '{REQUIRED_HEADINGS[0]}', but got '{first}'."
        )

    heading_positions = {}
    for idx, line in enumerate(lines):
        s = line.strip()
        if s in REQUIRED_HEADINGS:
            if s in heading_positions:
                raise ValueError(f"Duplicate section heading detected: {s}")
            heading_positions[s] = idx

    missing = [h for h in REQUIRED_HEADINGS if h not in heading_positions]
    if missing:
        raise ValueError(f"Missing required section(s): {', '.join(missing)}")

    positions = [heading_positions[h] for h in REQUIRED_HEADINGS]
    if positions != sorted(positions):
        raise ValueError("Section headings are not in the required order.")


def _safe_float(v: Any) -> float:
    try:
        return float(v or 0.0)
    except Exception:
        return 0.0


def _safe_int(v: Any) -> int:
    try:
        return int(v or 0)
    except Exception:
        return 0


def _safe_str(v: Any) -> str:
    return str(v or "").strip()


def _risk_label_to_detector_keys(risk_label: str) -> List[str]:
    s = _safe_str(risk_label).lower()
    if s == "structuring":
        return ["structured_deposits", "structured_payments"]
    if s == "pass-through":
        return ["pass_through"]
    if s == "layering":
        return ["layering"]
    if s == "round-figure amounts":
        return ["round_figures"]
    if s == "salary-like pattern":
        return ["salary_pattern"]
    if s == "cash-intensive activity":
        return ["cash_intensive"]
    if s == "third-party activity":
        return ["third_party"]
    if s == "recurrence":
        return ["recurrence"]
    return []


def _build_raw_row_map(raw_transactions: Any) -> Dict[int, Dict[str, Any]]:
    if raw_transactions is None or not hasattr(raw_transactions, "columns"):
        return {}

    df = raw_transactions
    if "ROW_ID" not in df.columns:
        try:
            df = df.copy()
            df["ROW_ID"] = df.index
        except Exception:
            return {}

    out: Dict[int, Dict[str, Any]] = {}
    for _, r in df.iterrows():
        rid = _safe_int(r.get("ROW_ID"))
        if rid < 0:
            continue

        date_s = _safe_str(r.get("DATE_STR") or r.get("DATE_RAW") or r.get("DATE"))
        dt = r.get("DATE")
        if hasattr(dt, "strftime"):
            try:
                date_s = dt.strftime("%Y-%m-%d")
            except Exception:
                pass

        credit = _safe_float(r.get("CREDIT"))
        debit = _safe_float(r.get("DEBIT"))
        amount = credit if credit > 0 else debit

        out[rid] = {
            "row_id": rid,
            "date": date_s,
            "transcode": _safe_str(r.get("TRANSCODE")),
            "description": _safe_str(r.get("DESCRIPTION_RAW") or r.get("DESCRIPTION")),
            "credit": credit,
            "debit": debit,
            "amount": amount,
            "balance": _safe_float(r.get("BALANCE")),
        }
    return out


def _pick_suspicious_rows_for_table(
    channel_code: str,
    risk_label: str,
    direction: str,
    suspicious_by_channel: Optional[Dict[str, Any]],
    raw_row_map: Dict[int, Dict[str, Any]],
) -> List[Dict[str, Any]]:
    by_channel = (suspicious_by_channel or {}).get("by_channel") or {}
    ch_block = by_channel.get(str(channel_code)) or {}
    rows = ch_block.get("rows") or []
    detector_keys = set(_risk_label_to_detector_keys(risk_label))
    if not rows or not detector_keys:
        return []

    picked: List[Dict[str, Any]] = []
    seen = set()

    for item in rows:
        if not isinstance(item, dict):
            continue

        dets = {str(x) for x in (item.get("detectors") or []) if str(x).strip()}
        if not dets.intersection(detector_keys):
            continue

        rid = _safe_int(item.get("row_id"))
        base = raw_row_map.get(rid, {})
        credit = _safe_float(base.get("credit") or item.get("credit"))
        debit = _safe_float(base.get("debit") or item.get("debit"))

        if direction == "credit" and credit <= 0:
            continue
        if direction == "debit" and debit <= 0:
            continue

        row = {
            "row_id": rid,
            "date": _safe_str(base.get("date") or item.get("date")),
            "transcode": _safe_str(base.get("transcode") or channel_code),
            "description": _safe_str(base.get("description") or item.get("description_raw")),
            "amount": credit if direction == "credit" else debit,
            "balance": _safe_float(base.get("balance")),
        }

        key = (
            row["row_id"],
            row["date"],
            row["transcode"],
            row["description"],
            row["amount"],
            row["balance"],
        )
        if key in seen:
            continue
        seen.add(key)
        picked.append(row)

    picked.sort(key=lambda x: (x.get("date") or "", x.get("row_id") or 0))
    return picked


def _add_suspicious_transactions_table(doc: Document, rows: List[Dict[str, Any]]) -> None:
    if not rows:
        return

    headers = ["Date", "Transcode", "Transaction Description", "Amount", "Balance"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        if hdr[i].paragraphs and hdr[i].paragraphs[0].runs:
            hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    total_amount = 0.0
    for r in rows:
        cells = table.add_row().cells
        amount = _safe_float(r.get("amount"))
        balance = _safe_float(r.get("balance"))
        total_amount += amount

        cells[0].text = _safe_str(r.get("date"))
        cells[1].text = _safe_str(r.get("transcode"))
        cells[2].text = _safe_str(r.get("description"))
        cells[3].text = f"K{amount:,.2f}" if abs(amount) > 1e-9 else ""
        cells[4].text = f"K{balance:,.2f}" if abs(balance) > 1e-9 else ""

    total_row = table.add_row().cells
    total_row[0].text = "TOTAL"
    total_row[1].text = ""
    total_row[2].text = f"{len(rows)} txn(s)"
    total_row[3].text = f"K{total_amount:,.2f}"
    total_row[4].text = ""

    for idx in (0, 2, 3):
        if total_row[idx].paragraphs and total_row[idx].paragraphs[0].runs:
            for run in total_row[idx].paragraphs[0].runs:
                run.bold = True


def _parse_channel_line(line: str) -> Tuple[str, str]:
    body = line[2:].strip()
    if " - " not in body:
        return "", ""
    first, rest = body.split(" - ", 1)
    code = _safe_str(first)
    desc = _safe_str(rest.split(":", 1)[0])
    return code, desc


def _looks_like_channel_header(line: str) -> bool:
    if not line.startswith("- "):
        return False
    body = line[2:].strip()
    return " - " in body and ":" in body


def _maybe_insert_risk_table(
    doc: Document,
    current_channel_code: str,
    current_direction: str,
    pending_risk_label: str,
    suspicious_by_channel: Optional[Dict[str, Any]],
    raw_row_map: Dict[int, Dict[str, Any]],
    inserted_keys: set,
) -> None:
    if not current_channel_code or not current_direction or not pending_risk_label:
        return

    key = (current_channel_code, current_direction, pending_risk_label)
    if key in inserted_keys:
        return

    table_rows = _pick_suspicious_rows_for_table(
        channel_code=current_channel_code,
        risk_label=pending_risk_label,
        direction=current_direction,
        suspicious_by_channel=suspicious_by_channel,
        raw_row_map=raw_row_map,
    )
    if table_rows:
        _add_suspicious_transactions_table(doc, table_rows)
        inserted_keys.add(key)




def _flush_pending_risk_table(
    doc: Document,
    current_channel_code: str,
    current_direction: str,
    pending_risk_label: str,
    suspicious_by_channel: Optional[Dict[str, Any]],
    raw_row_map: Dict[int, Dict[str, Any]],
    inserted_keys: set,
) -> None:
    _maybe_insert_risk_table(
        doc=doc,
        current_channel_code=current_channel_code,
        current_direction=current_direction,
        pending_risk_label=pending_risk_label,
        suspicious_by_channel=suspicious_by_channel,
        raw_row_map=raw_row_map,
        inserted_keys=inserted_keys,
    )


def generate_review_doc(
    client: Dict[str, Any],
    trigger: Dict[str, Any],
    aml_narrative: str,
    pivot_summary: List[Dict[str, Any]],
    material_channels: Optional[List[Dict[str, Any]]] = None,
    suspicious_by_channel: Optional[Dict[str, Any]] = None,
    raw_transactions: Any = None,
) -> str:
    validate_aml_narrative(aml_narrative)

    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_heading("", level=0)
    run = title.add_run("AML / CTF Statement Review Report")
    run.font.size = Pt(16)
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)

    doc.add_heading("Client Information", level=1)
    for k, v in (client or {}).items():
        p = doc.add_paragraph()
        p.add_run(f"{k}: ").bold = True
        p.add_run(str(v))
        set_paragraph_style(p)

    doc.add_heading("Trigger Information", level=1)
    for k, v in (trigger or {}).items():
        p = doc.add_paragraph()
        p.add_run(f"{k}: ").bold = True
        p.add_run(str(v))
        set_paragraph_style(p)

    doc.add_heading("Transaction Summary (Option A)", level=1)

    headers = [
        "Transaction Description",
        "Transaction Code",
        "DEPOSIT",
        "WITHDRAWAL",
        "Transaction Count",
        "CR%",
        "DR%",
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        if hdr_cells[i].paragraphs and hdr_cells[i].paragraphs[0].runs:
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in pivot_summary or []:
        cells = table.add_row().cells
        desc = row.get("DESCRIPTION", "")
        code = row.get("TRANSCODE", "")
        deposit = _safe_float(row.get("deposit", 0.0))
        withdrawal = _safe_float(row.get("withdrawal", 0.0))
        count = row.get("count", 0) or 0
        cr = _safe_float(row.get("CR%", 0.0))
        dr = _safe_float(row.get("DR%", 0.0))

        cells[0].text = str(desc)
        cells[1].text = str(code)
        cells[2].text = f"K{deposit:,.2f}" if abs(deposit) > 1e-9 else ""
        cells[3].text = f"K{withdrawal:,.2f}" if abs(withdrawal) > 1e-9 else ""
        cells[4].text = str(int(count))
        cells[5].text = f"{cr:.2f}%" if abs(cr) > 1e-9 else ""
        cells[6].text = f"{dr:.2f}%" if abs(dr) > 1e-9 else ""

    doc.add_page_break()
    doc.add_heading("Deterministic AML/CTF Assessment (No AI)", level=1)

    raw_lines = aml_narrative.split("\n")
    lines = [ln.rstrip() for ln in raw_lines if ln.strip()]
    required_set = set(REQUIRED_HEADINGS)
    raw_row_map = _build_raw_row_map(raw_transactions)

    header_indent = Inches(0.30)
    risk_indent = Inches(0.40)
    item_indent = Inches(0.48)
    item_indent_deep = Inches(0.58)
    item_indent_party = Inches(0.68)

    current_section = ""
    current_channel_code = ""
    current_direction = ""
    pending_risk_label = ""
    inserted_risk_tables = set()

    for line in lines:
        s = line.strip()

        if s in required_set:
            _flush_pending_risk_table(
                doc=doc,
                current_channel_code=current_channel_code,
                current_direction=current_direction,
                pending_risk_label=pending_risk_label,
                suspicious_by_channel=suspicious_by_channel,
                raw_row_map=raw_row_map,
                inserted_keys=inserted_risk_tables,
            )
            current_section = s
            current_channel_code = ""
            current_direction = ""
            pending_risk_label = ""
            hd = doc.add_heading(s, level=2)
            hd.paragraph_format.space_before = Pt(10)
            continue

        if line.startswith("- "):
            _flush_pending_risk_table(
                doc=doc,
                current_channel_code=current_channel_code,
                current_direction=current_direction,
                pending_risk_label=pending_risk_label,
                suspicious_by_channel=suspicious_by_channel,
                raw_row_map=raw_row_map,
                inserted_keys=inserted_risk_tables,
            )
            p = doc.add_paragraph(line[2:], style="List Bullet")
            set_paragraph_style(p)

            if current_section == REQUIRED_HEADINGS[0] and _looks_like_channel_header(line):
                current_direction = "credit"
                current_channel_code, _ = _parse_channel_line(line)
            elif current_section == REQUIRED_HEADINGS[1] and _looks_like_channel_header(line):
                current_direction = "debit"
                current_channel_code, _ = _parse_channel_line(line)
            else:
                current_channel_code = ""
                current_direction = ""

            pending_risk_label = ""
            continue

        if line.startswith("  - "):
            _flush_pending_risk_table(
                doc=doc,
                current_channel_code=current_channel_code,
                current_direction=current_direction,
                pending_risk_label=pending_risk_label,
                suspicious_by_channel=suspicious_by_channel,
                raw_row_map=raw_row_map,
                inserted_keys=inserted_risk_tables,
            )
            p = doc.add_paragraph(line[4:], style="List Bullet 2")
            p.paragraph_format.left_indent = header_indent
            p.paragraph_format.first_line_indent = Inches(-0.10)
            set_paragraph_style(p)
            pending_risk_label = ""
            continue

        if line.startswith("    - "):
            _flush_pending_risk_table(
                doc=doc,
                current_channel_code=current_channel_code,
                current_direction=current_direction,
                pending_risk_label=pending_risk_label,
                suspicious_by_channel=suspicious_by_channel,
                raw_row_map=raw_row_map,
                inserted_keys=inserted_risk_tables,
            )
            p = doc.add_paragraph(line[6:], style="List Bullet 3")
            p.paragraph_format.left_indent = risk_indent
            set_paragraph_style(p)
            pending_risk_label = line[6:].strip()
            continue

        if line.startswith("      ✓ "):
            p = doc.add_paragraph(line[6:])
            p.paragraph_format.left_indent = item_indent_deep
            p.paragraph_format.first_line_indent = Inches(0.0)
            set_paragraph_style(p)

            continue

        if line.startswith("      • "):
            p = doc.add_paragraph(line[6:])
            p.paragraph_format.left_indent = item_indent_party
            p.paragraph_format.first_line_indent = Inches(0.0)
            set_paragraph_style(p)
            continue

        if line.startswith("      "):
            p = doc.add_paragraph(line[6:])
            p.paragraph_format.left_indent = item_indent_deep
            p.paragraph_format.first_line_indent = Inches(0.0)
            set_paragraph_style(p)
            continue

        if line.startswith("    ✓ "):
            p = doc.add_paragraph(line[4:])
            p.paragraph_format.left_indent = item_indent
            p.paragraph_format.first_line_indent = Inches(0.0)
            set_paragraph_style(p)
            pending_risk_label = ""
            continue

        if len(s) > 2 and s[0].isdigit() and s[1] == ".":
            _flush_pending_risk_table(
                doc=doc,
                current_channel_code=current_channel_code,
                current_direction=current_direction,
                pending_risk_label=pending_risk_label,
                suspicious_by_channel=suspicious_by_channel,
                raw_row_map=raw_row_map,
                inserted_keys=inserted_risk_tables,
            )
            p = doc.add_paragraph(s, style="List Number")
            set_paragraph_style(p)
            pending_risk_label = ""
            continue

        _flush_pending_risk_table(
            doc=doc,
            current_channel_code=current_channel_code,
            current_direction=current_direction,
            pending_risk_label=pending_risk_label,
            suspicious_by_channel=suspicious_by_channel,
            raw_row_map=raw_row_map,
            inserted_keys=inserted_risk_tables,
        )
        p = doc.add_paragraph(s)
        set_paragraph_style(p)
        pending_risk_label = ""

    _flush_pending_risk_table(
        doc=doc,
        current_channel_code=current_channel_code,
        current_direction=current_direction,
        pending_risk_label=pending_risk_label,
        suspicious_by_channel=suspicious_by_channel,
        raw_row_map=raw_row_map,
        inserted_keys=inserted_risk_tables,
    )

    output = "AML_Review_Report.docx"
    doc.save(output)
    return output